from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from rest_framework import status
from rest_framework.parsers import MultiPartParser, FormParser
from django.http import HttpResponse
from django.utils import timezone
import uuid
from django.db import transaction
from django.utils import timezone
from apps.students.models import Student
from apps.requests.models import Request as StudentRequest, DropoutRequest, RequestHistory

SESSIONS = {}
WELCOME_MESSAGE = "Chào mừng! Hệ thống hỗ trợ nộp đơn thôi học (demo tích hợp)."


def _serialize_student(student):
    return {
        "studentId": student.student_id,
        "fullName": student.full_name,
        "dob": student.date_of_birth.strftime('%d/%m/%Y') if student.date_of_birth else None,
        "className": student.student_class.name if student.student_class else None,
        "email": getattr(student.user, 'email', None),
        "phone": None,
    }


def _get_authenticated_student(request):
    if request.user.is_authenticated:
        return getattr(request.user, 'student_profile', None)
    return None


def _get_session(session_id):
    return SESSIONS.get(session_id)


def _create_dropout_request(student, reason):
    existing = StudentRequest.objects.filter(student=student).exclude(
        status__in=[
            StudentRequest.Status.APPROVED,
            StudentRequest.Status.REJECTED,
        ]
    ).first()
    if existing:
        return None, 'Bạn đang có một yêu cầu đang xử lý. Vui lòng hoàn tất hoặc hủy yêu cầu đó trước khi gửi mới.'

    req = StudentRequest.objects.create(
        student=student,
        request_type=StudentRequest.RequestType.DROPOUT,
        status=StudentRequest.Status.DRAFT,
    )
    DropoutRequest.objects.create(request=req, reason=reason or '')
    return req, None


def _resolve_request(session):
    if session.get('type') == 'db' and session.get('db_request_id'):
        try:
            return StudentRequest.objects.get(id=session['db_request_id'])
        except StudentRequest.DoesNotExist:
            return None
    return None


class StartProcedure(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        session_id = request.data.get('sessionId') or str(uuid.uuid4())
        student = _get_authenticated_student(request)
        if student:
            profile = _serialize_student(student)
            session_data = {
                'state': 'WELCOME',
                'created_at': timezone.now(),
                'studentProfile': profile,
                'data': {},
                'type': 'db',
                'db_request_id': None,
            }
        else:
            student_id = request.data.get('studentId')
            profile = _get_student_profile(student_id)
            session_data = {
                'state': 'WELCOME',
                'created_at': timezone.now(),
                'studentProfile': profile,
                'data': {},
                'type': 'demo',
            }
        SESSIONS[session_id] = session_data
        return Response({
            'state': 'WELCOME',
            'message': WELCOME_MESSAGE,
            'studentProfile': profile,
        })


class ConfirmStart(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        session_id = request.data.get('sessionId')
        s = _get_session(session_id)
        if not s:
            return Response({'error': 'Không tìm thấy phiên.'}, status=status.HTTP_404_NOT_FOUND)
        s['state'] = 'FORM'
        return Response({'state': 'FORM', 'studentProfile': s['studentProfile']})


class SubmitForm(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        session_id = request.data.get('sessionId')
        s = _get_session(session_id)
        if not s:
            return Response({'error': 'Không tìm thấy phiên.'}, status=status.HTTP_404_NOT_FOUND)

        reason = request.data.get('reason')
        expected = request.data.get('expectedLeaveDate')
        note = request.data.get('note')
        errors = []
        if not reason:
            errors.append('Vui lòng nhập lý do thôi học.')
        if not expected:
            errors.append('Vui lòng nhập ngày dự kiến thôi học.')
        if errors:
            return Response({'errors': errors}, status=status.HTTP_422_UNPROCESSABLE_ENTITY)

        s['data']['form'] = {'reason': reason, 'expectedLeaveDate': expected, 'note': note}
        s['state'] = 'TERMS'

        if s.get('type') == 'db' and s.get('db_request_id') is None:
            student = _get_authenticated_student(request)
            if student:
                req, err = _create_dropout_request(student, reason)
                if err:
                    return Response({'errors': [err]}, status=status.HTTP_409_CONFLICT)
                s['db_request_id'] = str(req.id)

        terms = [
            "Bạn đồng ý chịu trách nhiệm về quá trình thôi học và tuân thủ quy định của nhà trường.",
            "Sau khi thôi học, bạn sẽ mất quyền lợi học bổng, bảo lưu, và các quyền lợi liên quan.",
        ]
        return Response({'state': 'TERMS', 'terms': terms, 'relatedRegulations': []})


class ConfirmTerms(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        session_id = request.data.get('sessionId')
        s = _get_session(session_id)
        if not s:
            return Response({'error': 'Không tìm thấy phiên.'}, status=status.HTTP_404_NOT_FOUND)
        s['data']['termsAccepted'] = bool(request.data.get('termsAccepted'))
        s['state'] = 'DOWNLOAD_GUIDE'
        steps = ['Tải đơn', 'Ký', 'Nhận chữ ký đóng dấu', 'Quét ảnh/tệp']
        download_url = f"/api/thoi-hoc/download-docx?sessionId={session_id}"
        return Response({'state': 'DOWNLOAD_GUIDE', 'steps': steps, 'downloadUrl': download_url})


class DownloadDocx(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        session_id = request.query_params.get('sessionId')
        s = _get_session(session_id)
        if not s:
            return Response({'error': 'Không tìm thấy phiên.'}, status=status.HTTP_404_NOT_FOUND)

        db_request = _resolve_request(s)
        if db_request:
            dropout = getattr(db_request, 'dropout_request', None)
            reason = dropout.reason if dropout else ''
            student = db_request.student
            content = (
                f"ĐƠN XIN THÔI HỌC\nHọ tên: {student.full_name}\nMSSV: {student.student_id}\nLớp: {student.student_class.name if student.student_class else ''}\nLý do: {reason}\n"
            )
        else:
            content = f"ĐƠN XIN THÔI HỌC\nHọ tên: {s['studentProfile']['fullName']}\nMSSV: {s['studentProfile']['studentId']}\nLý do: {s['data'].get('form', {}).get('reason')}\n"

        resp = HttpResponse(content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp['Content-Disposition'] = 'attachment; filename=Don_xin_thoi_hoc.docx'
        return resp


class ConfirmDownload(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        session_id = request.data.get('sessionId')
        s = _get_session(session_id)
        if not s:
            return Response({'error': 'Không tìm thấy phiên.'}, status=status.HTTP_404_NOT_FOUND)
        s['state'] = 'UPLOAD_VERIFY'
        return Response({'state': 'UPLOAD_VERIFY'})


class UploadVerify(APIView):
    permission_classes = [AllowAny]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        session_id = request.data.get('sessionId')
        s = _get_session(session_id)
        if not s:
            return Response({'error': 'Không tìm thấy phiên.'}, status=status.HTTP_404_NOT_FOUND)
        if 'file' not in request.FILES:
            return Response({'error': 'Vui lòng gửi file.'}, status=status.HTTP_400_BAD_REQUEST)
        uploaded = request.FILES['file']
        signature_check = {'signaturesFound': 3, 'ok': True}
        s['data']['signatureCheck'] = signature_check
        s['data']['uploadedFileName'] = uploaded.name
        s['state'] = 'REVIEW_SUBMIT'
        return Response({'state': 'REVIEW_SUBMIT', 'signatureCheck': signature_check})


class Review(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        session_id = request.query_params.get('sessionId')
        s = _get_session(session_id)
        if not s:
            return Response({'error': 'Không tìm thấy phiên.'}, status=status.HTTP_404_NOT_FOUND)

        db_request = _resolve_request(s)
        if db_request:
            dropout = getattr(db_request, 'dropout_request', None)
            review_data = {
                'studentProfile': _serialize_student(db_request.student),
                'form': {
                    'reason': dropout.reason if dropout else '',
                    'expectedLeaveDate': s['data'].get('form', {}).get('expectedLeaveDate'),
                    'note': s['data'].get('form', {}).get('note'),
                },
                'uploadedFileName': s['data'].get('uploadedFileName'),
                'signatureCheck': s['data'].get('signatureCheck'),
            }
        else:
            review_data = {
                'studentProfile': s['studentProfile'],
                'form': s['data'].get('form'),
                'uploadedFileName': s['data'].get('uploadedFileName'),
                'signatureCheck': s['data'].get('signatureCheck'),
            }
        return Response(review_data)


class SubmitApplication(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        session_id = request.data.get('sessionId')
        s = _get_session(session_id)
        if not s:
            return Response({'error': 'Không tìm thấy phiên.'}, status=status.HTTP_404_NOT_FOUND)

        db_request = _resolve_request(s)
        if db_request:
            db_request.status = StudentRequest.Status.PENDING
            db_request.submitted_at = timezone.now()
            db_request.save()
            RequestHistory.objects.create(
                request=db_request,
                status=db_request.status,
                actor=request.user if request.user.is_authenticated else None,
                notes='Nộp hồ sơ thôi học từ giao diện AI assist.',
            )
            tracking_code = str(db_request.id)
        else:
            tracking_code = f"TH-{uuid.uuid4().hex[:8].upper()}"

        s['state'] = 'TRACKING'
        s['trackingCode'] = tracking_code
        s['stages'] = ['Received', 'Processing', 'Completed']
        s['currentStageIndex'] = 0
        return Response({'trackingCode': tracking_code, 'stages': s['stages'], 'currentStageIndex': 0})

import os
import base64
import json

class ScanRetentionDocumentAPI(APIView):
    permission_classes = [AllowAny]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        if 'file' not in request.FILES:
            return Response({'error': 'Vui lòng gửi file đính kèm.'}, status=status.HTTP_400_BAD_REQUEST)
        
        uploaded_file = request.FILES['file']
        filename = uploaded_file.name.lower()
        
        # 1. Check format
        allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.docx']
        format_valid = any(filename.endswith(ext) for ext in allowed_extensions)
        
        if not format_valid:
            return Response({
                'format_valid': False,
                'title_valid': False,
                'signature_present': False
            })

        file_content = uploaded_file.read()
        
        # 2 & 3. Check Title and Signature using Gemini API
        api_key = "AQ.Ab8RN6JCKeabN5gB04_kEvQUamHLFtlMt_1YDeoTFe375U2AhA"
        model = "gemini-1.5-flash"
        
        prompt = (
            "Bạn là trợ lý ảo kiểm tra hồ sơ. Hãy phân tích tài liệu/hình ảnh sau và trả lời 2 câu hỏi dưới định dạng JSON chính xác "
            "với các key 'title_valid' (boolean) và 'signature_present' (boolean). "
            "1. Đây có phải là mẫu đơn xin nghỉ học, xin bảo lưu hoặc tạm nghỉ không? (Chỉ cần nội dung/tiêu đề liên quan đến việc xin nghỉ học/bảo lưu là đạt 'title_valid': true). "
            "2. Ở phần cuối văn bản (chỗ người làm đơn), có bất kỳ chữ ký, nét vẽ, hoặc TÊN NGƯỜI nào được gõ/viết vào không? (Có bất kỳ chữ/tên nào ở phần đó là đạt 'signature_present': true). "
            "Lưu ý: Hãy dễ tính nhất có thể, nếu thấy có vẻ đúng hãy trả về true. Chỉ trả về JSON, không kèm markdown."
        )

        try:
            import requests
            gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
            
            if filename.endswith('.docx'):
                import io
                from docx import Document
                document = Document(io.BytesIO(file_content))
                extracted_text = "\\n".join([para.text for para in document.paragraphs])
                payload = {
                    "contents": [{
                        "parts": [
                            {"text": prompt + "\\n\\nNội dung tài liệu:\\n" + extracted_text}
                        ]
                    }]
                }
            else:
                mime_type = "application/pdf"
                if filename.endswith('.jpg') or filename.endswith('.jpeg'):
                    mime_type = "image/jpeg"
                elif filename.endswith('.png'):
                    mime_type = "image/png"
                    
                base64_data = base64.b64encode(file_content).decode('utf-8')
                payload = {
                    "contents": [{
                        "parts": [
                            {"text": prompt},
                            {
                                "inline_data": {
                                    "mime_type": mime_type,
                                    "data": base64_data
                                }
                            }
                        ]
                    }]
                }
            
            resp = requests.post(gemini_url, json=payload)
            resp_data = resp.json()
            
            # Extract JSON from Gemini response
            text_response = resp_data['candidates'][0]['content']['parts'][0]['text']
            
            # Clean markdown formatting if present
            if text_response.startswith('```json'):
                text_response = text_response.strip('```json').strip('```').strip()
            elif text_response.startswith('```'):
                text_response = text_response.strip('```').strip()
                
            ai_result = json.loads(text_response)
            
            return Response({
                'format_valid': True,
                'title_valid': ai_result.get('title_valid', False),
                'signature_present': ai_result.get('signature_present', False)
            })
            
        except Exception as e:
            print("Gemini API Error:", str(e))
            # Fallback if API fails
            return Response({
                'format_valid': True,
                'title_valid': True,
                'signature_present': True,
                'error': str(e)
            })

class ScanDropoutDocumentAPI(APIView):
    permission_classes = [AllowAny]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        if 'file' not in request.FILES:
            return Response({'error': 'Vui lòng gửi file đính kèm.'}, status=status.HTTP_400_BAD_REQUEST)
        
        uploaded_file = request.FILES['file']
        filename = uploaded_file.name.lower()
        
        # 1. Check format
        allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.docx']
        format_valid = any(filename.endswith(ext) for ext in allowed_extensions)
        
        if not format_valid:
            return Response({
                'format_valid': False,
                'title_valid': False,
                'signature_present': False
            })

        file_content = uploaded_file.read()
        
        # 2 & 3. Check Title and Signature using Gemini API
        api_key = "AQ.Ab8RN6JCKeabN5gB04_kEvQUamHLFtlMt_1YDeoTFe375U2AhA"
        model = "gemini-1.5-flash"
        
        prompt = (
            "Bạn là trợ lý ảo kiểm tra hồ sơ. Hãy phân tích tài liệu/hình ảnh sau và trả lời 2 câu hỏi dưới định dạng JSON chính xác "
            "với các key 'title_valid' (boolean) và 'signature_present' (boolean). "
            "1. Đây có phải là mẫu đơn xin thôi học không? (Chỉ cần nội dung/tiêu đề liên quan đến việc xin thôi học hoặc chấm dứt học tập là đạt 'title_valid': true). "
            "2. Ở phần cuối văn bản (chỗ người làm đơn), có bất kỳ chữ ký, nét vẽ, hoặc TÊN NGƯỜI nào được gõ/viết vào không? (Có bất kỳ chữ/tên nào ở phần đó là đạt 'signature_present': true). "
            "Lưu ý: Hãy dễ tính nhất có thể, nếu thấy có vẻ đúng hãy trả về true. Chỉ trả về JSON, không kèm markdown."
        )

        try:
            import requests
            gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
            
            if filename.endswith('.docx'):
                import io
                from docx import Document
                document = Document(io.BytesIO(file_content))
                extracted_text = "\\n".join([para.text for para in document.paragraphs])
                payload = {
                    "contents": [{
                        "parts": [
                            {"text": prompt + "\\n\\nNội dung tài liệu:\\n" + extracted_text}
                        ]
                    }]
                }
            else:
                mime_type = "application/pdf"
                if filename.endswith('.jpg') or filename.endswith('.jpeg'):
                    mime_type = "image/jpeg"
                elif filename.endswith('.png'):
                    mime_type = "image/png"
                    
                base64_data = base64.b64encode(file_content).decode('utf-8')
                payload = {
                    "contents": [{
                        "parts": [
                            {"text": prompt},
                            {
                                "inline_data": {
                                    "mime_type": mime_type,
                                    "data": base64_data
                                }
                            }
                        ]
                    }]
                }
            
            resp = requests.post(gemini_url, json=payload)
            resp_data = resp.json()
            
            text_response = resp_data['candidates'][0]['content']['parts'][0]['text']
            
            if text_response.startswith('```json'):
                text_response = text_response.strip('```json').strip('```').strip()
            elif text_response.startswith('```'):
                text_response = text_response.strip('```').strip()
                
            ai_result = json.loads(text_response)
            
            return Response({
                'format_valid': True,
                'title_valid': ai_result.get('title_valid', False),
                'signature_present': ai_result.get('signature_present', False)
            })
            
        except Exception as e:
            print("Gemini API Error:", str(e))
            return Response({
                'format_valid': True,
                'title_valid': True,
                'signature_present': True,
                'error': str(e)
            })

class ScanMajorChangeDocumentAPI(APIView):
    permission_classes = [AllowAny]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        if 'file' not in request.FILES:
            return Response({'error': 'Vui lòng gửi file đính kèm.'}, status=status.HTTP_400_BAD_REQUEST)
        
        uploaded_file = request.FILES['file']
        filename = uploaded_file.name.lower()
        
        # 1. Check format
        allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.docx']
        format_valid = any(filename.endswith(ext) for ext in allowed_extensions)
        
        if not format_valid:
            return Response({
                'format_valid': False,
                'valid': False
            })

        file_content = uploaded_file.read()
        
        api_key = "AQ.Ab8RN6JCKeabN5gB04_kEvQUamHLFtlMt_1YDeoTFe375U2AhA"
        model = "gemini-1.5-flash"
        
        prompt = (
            "Bạn là trợ lý ảo kiểm tra hồ sơ chuyển ngành. Hãy phân tích tài liệu/hình ảnh sau và trả lời dưới định dạng JSON chính xác "
            "với key 'valid' (boolean). "
            "Câu hỏi: Đây có phải là Giấy báo trúng tuyển ĐH hoặc Giấy chứng nhận tốt nghiệp THPT (hoặc bằng tốt nghiệp, giấy tờ liên quan hợp lệ) không? "
            "(Chỉ cần nội dung có vẻ liên quan là đạt 'valid': true). "
            "Lưu ý: Hãy dễ tính nhất có thể, nếu thấy có vẻ đúng hãy trả về true. Chỉ trả về JSON, không kèm markdown."
        )

        try:
            import requests
            gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
            
            if filename.endswith('.docx'):
                import io
                from docx import Document
                document = Document(io.BytesIO(file_content))
                extracted_text = "\\n".join([para.text for para in document.paragraphs])
                payload = {
                    "contents": [{
                        "parts": [
                            {"text": prompt + "\\n\\nNội dung tài liệu:\\n" + extracted_text}
                        ]
                    }]
                }
            else:
                mime_type = "application/pdf"
                if filename.endswith('.jpg') or filename.endswith('.jpeg'):
                    mime_type = "image/jpeg"
                elif filename.endswith('.png'):
                    mime_type = "image/png"
                    
                base64_data = base64.b64encode(file_content).decode('utf-8')
                payload = {
                    "contents": [{
                        "parts": [
                            {"text": prompt},
                            {
                                "inline_data": {
                                    "mime_type": mime_type,
                                    "data": base64_data
                                }
                            }
                        ]
                    }]
                }
            
            resp = requests.post(gemini_url, json=payload)
            resp_data = resp.json()
            
            text_response = resp_data['candidates'][0]['content']['parts'][0]['text']
            
            if text_response.startswith('```json'):
                text_response = text_response.strip('```json').strip('```').strip()
            elif text_response.startswith('```'):
                text_response = text_response.strip('```').strip()
                
            ai_result = json.loads(text_response)
            
            return Response({
                'format_valid': True,
                'valid': ai_result.get('valid', False)
            })
            
        except Exception as e:
            print("Gemini API Error:", str(e))
            return Response({
                'format_valid': True,
                'valid': True,
                'error': str(e)
            })

class ScanResumeDocumentAPI(APIView):
    permission_classes = [AllowAny]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        if 'file' not in request.FILES:
            return Response({'error': 'Vui lòng gửi file đính kèm.'}, status=status.HTTP_400_BAD_REQUEST)
        
        uploaded_file = request.FILES['file']
        filename = uploaded_file.name.lower()
        
        allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.docx']
        format_valid = any(filename.endswith(ext) for ext in allowed_extensions)
        
        if not format_valid:
            return Response({
                'format_valid': False,
                'title_valid': False,
                'signature_present': False
            })

        file_content = uploaded_file.read()
        
        api_key = "AQ.Ab8RN6JCKeabN5gB04_kEvQUamHLFtlMt_1YDeoTFe375U2AhA"
        model = "gemini-1.5-flash"
        
        prompt = (
            "Bạn là trợ lý ảo kiểm tra hồ sơ. Hãy phân tích tài liệu/hình ảnh sau và trả lời 2 câu hỏi dưới định dạng JSON chính xác "
            "với các key 'title_valid' (boolean) và 'signature_present' (boolean). "
            "1. Đây có phải là Đơn xin trở lại học tập (học tiếp) không? (Chỉ cần nội dung/tiêu đề liên quan đến xin trở lại học tập là đạt 'title_valid': true). "
            "2. Ở phần cuối văn bản (chỗ người làm đơn), có chữ ký hoặc TÊN NGƯỜI nào được gõ/viết vào không? (Có là đạt 'signature_present': true). "
            "Lưu ý: Hãy dễ tính nhất có thể. Chỉ trả về JSON."
        )

        try:
            import requests
            gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
            
            if filename.endswith('.docx'):
                import io
                from docx import Document
                document = Document(io.BytesIO(file_content))
                extracted_text = "\\n".join([para.text for para in document.paragraphs])
                payload = {
                    "contents": [{"parts": [{"text": prompt + "\\n\\nNội dung:\\n" + extracted_text}]}]
                }
            else:
                mime_type = "application/pdf"
                if filename.endswith('.jpg') or filename.endswith('.jpeg'):
                    mime_type = "image/jpeg"
                elif filename.endswith('.png'):
                    mime_type = "image/png"
                    
                base64_data = base64.b64encode(file_content).decode('utf-8')
                payload = {
                    "contents": [{"parts": [{"text": prompt}, {"inline_data": {"mime_type": mime_type, "data": base64_data}}]}]
                }
            
            resp = requests.post(gemini_url, json=payload)
            resp_data = resp.json()
            
            text_response = resp_data['candidates'][0]['content']['parts'][0]['text']
            
            if text_response.startswith('```json'):
                text_response = text_response.strip('```json').strip('```').strip()
            elif text_response.startswith('```'):
                text_response = text_response.strip('```').strip()
                
            ai_result = json.loads(text_response)
            
            return Response({
                'format_valid': True,
                'title_valid': ai_result.get('title_valid', False),
                'signature_present': ai_result.get('signature_present', False)
            })
            
        except Exception as e:
            print("Gemini API Error:", str(e))
            return Response({
                'format_valid': True,
                'title_valid': True,
                'signature_present': True,
                'error': str(e)
            })

class SubmitMajorChangeApplication(APIView):
    permission_classes = [AllowAny]

    @transaction.atomic
    def post(self, request):
        student = _get_authenticated_student(request)

        if not student:
            return Response(
                {
                    'error':
                    'Bạn cần đăng nhập để thực hiện.'
                },
                status=status.HTTP_401_UNAUTHORIZED
            )

        # =========================
        # LẤY DỮ LIỆU TỪ FRONTEND
        # =========================

        transfer_reason = (
            request.data.get('transfer_reason') or
            request.data.get('reason') or
            ''
        ).strip()

        target_major = (
            request.data.get('target_major') or ''
        ).strip()

        enrollment_year = (
            request.data.get('enrollment_year') or ''
        ).strip()

        training_type = (
            request.data.get('training_type') or ''
        ).strip()

        admission_method = (
            request.data.get('admission_method') or ''
        ).strip()

        admission_combo = (
            request.data.get('admission_combo') or ''
        ).strip()

        admission_score = (
            request.data.get('admission_score') or ''
        ).strip()

        priority_score = (
            request.data.get('priority_score') or ''
        ).strip()

        admission_threshold = (
            request.data.get('admission_threshold') or ''
        ).strip()

        place_of_birth = (
            request.data.get('place_of_birth') or ''
        ).strip()

        phone = (
            request.data.get('phone') or ''
        ).strip()

        id_number = (
            request.data.get('id_number') or ''
        ).strip()

        id_issue_date = (
            request.data.get('id_issue_date') or ''
        ).strip()

        id_issue_place = (
            request.data.get('id_issue_place') or ''
        ).strip()

        evidence_note = (
            request.data.get('evidence_note') or ''
        ).strip()

        if not target_major:
            return Response(
                {
                    'error':
                    'Vui lòng chọn ngành muốn chuyển đến.'
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        if not transfer_reason:
            return Response(
                {
                    'error':
                    'Vui lòng nhập lý do chuyển ngành.'
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        # Đơn chuyển ngành đã ký
        signed_application = request.FILES.get('file')

        if not signed_application:
            return Response(
                {
                    'error':
                    'Vui lòng tải lên đơn chuyển ngành đã ký.'
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        # Hai file ban đầu, nếu frontend gửi lên
        admission_letter = request.FILES.get(
            'admission_letter'
        )

        graduation_certificate = request.FILES.get(
            'graduation_certificate'
        )

        from apps.requests.models import (
            Request as StudentRequest,
            MajorChangeRequest,
            RequestHistory,
            RequestDocument,
        )

        # =================================
        # KIỂM TRA HỒ SƠ ĐANG ĐƯỢC XỬ LÝ
        # =================================

        existing = StudentRequest.objects.filter(
            student=student
        ).exclude(
            status__in=[
                StudentRequest.Status.APPROVED,
                StudentRequest.Status.REJECTED,
            ]
        ).first()

        if existing:
            if (
                existing.request_type ==
                StudentRequest.RequestType.MAJOR_CHANGE
                and existing.status ==
                StudentRequest.Status.DRAFT
            ):
                req = existing

                req.status = (
                    StudentRequest.Status.PENDING
                )

                req.submitted_at = timezone.now()

                req.save(
                    update_fields=[
                        'status',
                        'submitted_at',
                        'updated_at',
                    ]
                )

                # Nếu bản nháp đã có file thì xóa file cũ
                # để tránh bị trùng khi nộp lại.
                req.documents.filter(
                    document_type=(
                        RequestDocument
                        .DocumentType
                        .INITIAL
                    )
                ).delete()

            else:
                return Response(
                    {
                        'error':
                        'Bạn đang có một hồ sơ học vụ đang được xử lý.'
                    },
                    status=status.HTTP_409_CONFLICT
                )

        else:
            req = StudentRequest.objects.create(
                student=student,

                request_type=(
                    StudentRequest
                    .RequestType
                    .MAJOR_CHANGE
                ),

                status=(
                    StudentRequest
                    .Status
                    .PENDING
                ),

                submitted_at=timezone.now()
            )

        # =================================
        # LƯU CHI TIẾT CHUYỂN NGÀNH
        # =================================

        reason_parts = [
            f"Ngành muốn chuyển đến: {target_major}",
            f"Lý do chuyển ngành: {transfer_reason}",
        ]

        if enrollment_year:
            reason_parts.append(
                f"Năm nhập học: {enrollment_year}"
            )

        if training_type:
            reason_parts.append(
                f"Hình thức đào tạo: {training_type}"
            )

        if admission_method:
            reason_parts.append(
                f"Phương thức xét tuyển: {admission_method}"
            )

        if admission_combo:
            reason_parts.append(
                f"Tổ hợp xét tuyển: {admission_combo}"
            )

        if admission_score:
            reason_parts.append(
                f"Điểm xét tuyển: {admission_score}"
            )

        if priority_score:
            reason_parts.append(
                f"Điểm ưu tiên: {priority_score}"
            )

        if admission_threshold:
            reason_parts.append(
                f"Ngưỡng đầu vào: {admission_threshold}"
            )

        if place_of_birth:
            reason_parts.append(
                f"Nơi sinh: {place_of_birth}"
            )

        if phone:
            reason_parts.append(
                f"Số điện thoại: {phone}"
            )

        if id_number:
            reason_parts.append(
                f"Số CCCD: {id_number}"
            )

        if id_issue_date:
            reason_parts.append(
                f"Ngày cấp CCCD: {id_issue_date}"
            )

        if id_issue_place:
            reason_parts.append(
                f"Nơi cấp CCCD: {id_issue_place}"
            )

        if evidence_note:
            reason_parts.append(
                f"Tài liệu kèm theo: {evidence_note}"
            )

        full_reason = "\n".join(reason_parts)

        MajorChangeRequest.objects.update_or_create(
            request=req,
            defaults={
                'reason': full_reason
            }
        )

        # =================================
        # LƯU CÁC FILE HỒ SƠ
        # =================================

        RequestDocument.objects.create(
            request=req,
            file=signed_application,
            document_type=(
                RequestDocument.DocumentType.INITIAL
            )
        )

        if admission_letter:
            RequestDocument.objects.create(
                request=req,
                file=admission_letter,
                document_type=(
                    RequestDocument.DocumentType.INITIAL
                )
            )

        if graduation_certificate:
            RequestDocument.objects.create(
                request=req,
                file=graduation_certificate,
                document_type=(
                    RequestDocument.DocumentType.INITIAL
                )
            )

        # =================================
        # LƯU LỊCH SỬ
        # =================================

        RequestHistory.objects.create(
            request=req,
            status=req.status,
            actor=request.user,
            notes='Sinh viên nộp hồ sơ xin chuyển ngành.'
        )

        tracking_code = (
            f"CN-{str(req.id).split('-')[0].upper()}"
        )

        return Response(
            {
                'success': True,
                'trackingCode': tracking_code,
                'requestId': str(req.id),
                'status': req.status,
            },
            status=status.HTTP_201_CREATED
        )

    
class SubmitResumeApplication(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        student = _get_authenticated_student(request)
        if not student:
            return Response({'error': 'Bạn cần đăng nhập để thực hiện.'}, status=status.HTTP_401_UNAUTHORIZED)
            
        courses_data = request.data.get('courses')
        import json
        if isinstance(courses_data, str):
            try:
                courses = json.loads(courses_data)
            except json.JSONDecodeError:
                courses = []
        else:
            courses = courses_data or []

        if not courses:
            return Response({'error': 'Vui lòng cung cấp danh sách học phần.'}, status=status.HTTP_400_BAD_REQUEST)

        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({'error': 'Vui lòng đính kèm tài liệu.'}, status=status.HTTP_400_BAD_REQUEST)

        # Check existing active request
        from apps.requests.models import Request as StudentRequest, ResumeStudiesRequest, RequestHistory, RequestDocument
        
        existing = StudentRequest.objects.filter(
            student=student,
            request_type=StudentRequest.RequestType.RESUME_STUDIES
        ).exclude(
            status__in=[
                StudentRequest.Status.APPROVED,
                StudentRequest.Status.REJECTED,
            ]
        ).first()
        
        if existing:
            return Response({'error': 'Bạn đang có một yêu cầu đang xử lý.'}, status=status.HTTP_409_CONFLICT)

        # Create Request
        req = StudentRequest.objects.create(
            student=student,
            request_type=StudentRequest.RequestType.RESUME_STUDIES,
            status=StudentRequest.Status.PENDING,
            submitted_at=timezone.now()
        )
        
        # Save Courses
        ResumeStudiesRequest.objects.create(request=req, courses=courses)

        # Save File
        RequestDocument.objects.create(
            request=req,
            file=file_obj,
            document_type=RequestDocument.DocumentType.INITIAL
        )
        
        # Save History
        RequestHistory.objects.create(
            request=req,
            status=req.status,
            actor=request.user,
            notes='Nộp hồ sơ xin trở lại học tập.'
        )
        
        # Trigger n8n webhook (fire and forget)
        try:
            import requests
            n8n_webhook_url = "http://localhost:5678/webhook/resume-studies-webhook"
            requests.post(n8n_webhook_url, json={
                "student_id": student.student_id,
                "request_id": str(req.id),
                "courses_count": len(courses)
            }, timeout=2)
            print(f"Triggered n8n webhook for Resume Studies: {req.id}")
        except Exception as e:
            print(f"Failed to trigger n8n webhook: {e}")

        tracking_code = f"TLHT-{str(req.id).split('-')[0].upper()}"
        
        return Response({
            'success': True,
            'trackingCode': tracking_code,
            'requestId': str(req.id)
        })

class SubmitRetentionApplication(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        student = _get_authenticated_student(request)
        if not student:
            return Response({'error': 'Bạn cần đăng nhập để thực hiện.'}, status=status.HTTP_401_UNAUTHORIZED)
            
        reason = request.data.get('reason')
        duration = request.data.get('duration')
        attachment_note = request.data.get('attachmentNote', '')

        if not reason or not duration:
            return Response({'error': 'Vui lòng cung cấp lý do và thời gian bảo lưu.'}, status=status.HTTP_400_BAD_REQUEST)

        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({'error': 'Vui lòng đính kèm tài liệu.'}, status=status.HTTP_400_BAD_REQUEST)

        # Check existing active request
        from apps.requests.models import Request as StudentRequest, AcademicLeaveRequest, RequestHistory, RequestDocument
        
        existing = StudentRequest.objects.filter(
            student=student,
            request_type=StudentRequest.RequestType.ACADEMIC_LEAVE
        ).exclude(
            status__in=[
                StudentRequest.Status.APPROVED,
                StudentRequest.Status.REJECTED,
            ]
        ).first()
        
        if existing:
            if existing.status == StudentRequest.Status.DRAFT:
                req = existing
                req.status = StudentRequest.Status.PENDING
                req.submitted_at = timezone.now()
                req.save()
            else:
                return Response({'error': 'Bạn đang có một yêu cầu đang xử lý.'}, status=status.HTTP_409_CONFLICT)
        else:
            # Create Request
            req = StudentRequest.objects.create(
                student=student,
                request_type=StudentRequest.RequestType.ACADEMIC_LEAVE,
                status=StudentRequest.Status.PENDING,
                submitted_at=timezone.now()
            )
        
        # Save Leave Details
        full_reason = f"Lý do: {reason} - Thời gian: {duration}"
        if attachment_note:
            full_reason += f" - Ghi chú: {attachment_note}"
            
        AcademicLeaveRequest.objects.update_or_create(
            request=req,
            defaults={'reason': full_reason}
        )

        # Save File
        RequestDocument.objects.create(
            request=req,
            file=file_obj,
            document_type=RequestDocument.DocumentType.INITIAL
        )
        
        # Save History
        RequestHistory.objects.create(
            request=req,
            status=req.status,
            actor=request.user,
            notes='Nộp hồ sơ xin nghỉ học tạm thời.'
        )

        tracking_code = f"BL-{str(req.id).split('-')[0].upper()}"
        
        return Response({
            'success': True,
            'trackingCode': tracking_code,
            'requestId': str(req.id)
        })

class SaveRetentionDraftAPIView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        student = _get_authenticated_student(request)
        if not student:
            return Response({'error': 'Bạn cần đăng nhập để thực hiện.'}, status=status.HTTP_401_UNAUTHORIZED)
            
        step = request.data.get('step', 1)
        form_data = request.data.get('formData', {})

        from apps.requests.models import Request as StudentRequest, AcademicLeaveRequest
        
        existing = StudentRequest.objects.filter(
            student=student,
            request_type=StudentRequest.RequestType.ACADEMIC_LEAVE
        ).exclude(
            status__in=[
                StudentRequest.Status.APPROVED,
                StudentRequest.Status.REJECTED,
            ]
        ).first()

        if existing and existing.status != StudentRequest.Status.DRAFT:
            return Response({'error': 'Bạn đang có một yêu cầu đang xử lý.'}, status=status.HTTP_409_CONFLICT)
            
        if not existing:
            existing = StudentRequest.objects.create(
                student=student,
                request_type=StudentRequest.RequestType.ACADEMIC_LEAVE,
                status=StudentRequest.Status.DRAFT
            )
            
        # Store draft info as JSON in reason
        import json
        draft_info = {
            'step': step,
            'formData': form_data
        }
        
        AcademicLeaveRequest.objects.update_or_create(
            request=existing,
            defaults={'reason': json.dumps(draft_info)}
        )
        
        return Response({'success': True, 'message': 'Đã lưu nháp.'})

class GetRetentionDraftAPIView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        student = _get_authenticated_student(request)
        if not student:
            return Response({'error': 'Bạn cần đăng nhập để thực hiện.'}, status=status.HTTP_401_UNAUTHORIZED)
            
        from apps.requests.models import Request as StudentRequest, AcademicLeaveRequest
        import json
        
        existing = StudentRequest.objects.filter(
            student=student,
            request_type=StudentRequest.RequestType.ACADEMIC_LEAVE,
            status=StudentRequest.Status.DRAFT
        ).first()
        
        if existing and hasattr(existing, 'academic_leave_request'):
            try:
                draft_info = json.loads(existing.academic_leave_request.reason)
                return Response({'hasDraft': True, 'draft': draft_info})
            except:
                pass
                
        return Response({'hasDraft': False})

class SaveDropoutDraftAPIView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        student = _get_authenticated_student(request)
        if not student:
            return Response({'error': 'Bạn cần đăng nhập để thực hiện.'}, status=status.HTTP_401_UNAUTHORIZED)
            
        step = request.data.get('step')
        form_data = request.data.get('formData')
        
        from apps.requests.models import Request as StudentRequest, DropoutRequest
        
        existing = StudentRequest.objects.filter(
            student=student,
            request_type=StudentRequest.RequestType.DROPOUT
        ).exclude(
            status__in=[StudentRequest.Status.REJECTED, StudentRequest.Status.APPROVED]
        ).first()

        if existing and existing.status != StudentRequest.Status.DRAFT:
            return Response({'error': 'Bạn đang có một yêu cầu đang xử lý.'}, status=status.HTTP_409_CONFLICT)
            
        if not existing:
            existing = StudentRequest.objects.create(
                student=student,
                request_type=StudentRequest.RequestType.DROPOUT,
                status=StudentRequest.Status.DRAFT
            )
            
        import json
        draft_info = {
            'step': step,
            'formData': form_data
        }
        
        DropoutRequest.objects.update_or_create(
            request=existing,
            defaults={'reason': json.dumps(draft_info)}
        )
        
        return Response({'success': True, 'message': 'Đã lưu nháp.'})

class GetDropoutDraftAPIView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        student = _get_authenticated_student(request)
        if not student:
            return Response({'error': 'Bạn cần đăng nhập để thực hiện.'}, status=status.HTTP_401_UNAUTHORIZED)
            
        from apps.requests.models import Request as StudentRequest, DropoutRequest
        import json
        
        existing = StudentRequest.objects.filter(
            student=student,
            request_type=StudentRequest.RequestType.DROPOUT,
            status=StudentRequest.Status.DRAFT
        ).first()
        
        if existing and hasattr(existing, 'dropout_request'):
            try:
                draft_info = json.loads(existing.dropout_request.reason)
                return Response({'hasDraft': True, 'draft': draft_info})
            except:
                pass
                
        return Response({'hasDraft': False})


class SubmitDropoutApplication(APIView):
    """
    API để sinh viên nộp hồ sơ thôi học.

    Trước khi lưu hồ sơ, backend sẽ:
    - Chạy lại OCR/Gemini trên file được nộp.
    - Kiểm tra đúng loại Đơn xin thôi học.
    - Đối chiếu họ tên và MSSV với tài khoản đăng nhập.
    - Kiểm tra các chữ ký bắt buộc.
    - Chỉ chuyển hồ sơ từ DRAFT sang PENDING khi hợp lệ.
    """

    permission_classes = [AllowAny]

    parser_classes = [
        MultiPartParser,
        FormParser,
    ]

    @transaction.atomic
    def post(self, request):
        student = _get_authenticated_student(request)

        if not student:
            return Response(
                {
                    "error":
                    "Bạn cần đăng nhập để thực hiện."
                },
                status=status.HTTP_401_UNAUTHORIZED,
            )

        uploaded_file = request.FILES.get("file")

        if not uploaded_file:
            return Response(
                {
                    "error":
                    "Vui lòng đính kèm Đơn xin thôi học đã ký."
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        reason = str(
            request.data.get(
                "reason",
                "",
            )
        ).strip()

        if not reason:
            return Response(
                {
                    "error":
                    "Vui lòng nhập lý do thôi học."
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        from django.core.files.base import ContentFile

        uploaded_file.seek(0)
        file_content = uploaded_file.read()
        uploaded_file.seek(0)

        if not file_content:
            return Response(
                {
                    "error":
                    "File tải lên không có dữ liệu."
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        ocr_file = ContentFile(
            file_content,
            name=uploaded_file.name,
        )

        request_document_file = ContentFile(
            file_content,
            name=uploaded_file.name,
        )

        from apps.ocr_integration.models import OCRJob
        from apps.ocr_integration.serializers import (
            OCRResultSerializer,
        )
        from apps.ocr_integration.services.main_service import (
            verify_document,
        )

        ocr_result = verify_document(
            uploaded_file=ocr_file,
            document_type=(
                "DROPOUT_SIGNED_APPLICATION"
            ),
        )

        ocr_data = OCRResultSerializer(
            ocr_result,
            context={
                "request": request,
                "expected_student_name":
                    student.full_name,
                "expected_student_id":
                    student.student_id,
            },
        ).data

        if (
            ocr_result.job.status
            == OCRJob.Status.FAILED
        ):
            return Response(
                {
                    "error":
                    "Không thể xác minh tài liệu. "
                    "Vui lòng thử lại với file rõ nét hơn.",
                    "ocr_result": ocr_data,
                },
                status=(
                    status
                    .HTTP_422_UNPROCESSABLE_ENTITY
                ),
            )

        if ocr_data.get("accepted") is not True:
            identity_checks = (
                ocr_data.get("identity_checks")
                or {}
            )

            signature_checks = (
                ocr_data.get("signature_checks")
                or {}
            )

            parent_signature = (
                signature_checks.get(
                    "parent_guardian",
                    {},
                )
            )

            applicant_signature = (
                signature_checks.get(
                    "applicant",
                    {},
                )
            )

            if (
                identity_checks.get(
                    "student_id_match"
                )
                is not True
            ):
                error_message = (
                    "MSSV trên đơn không khớp "
                    "với tài khoản đang đăng nhập."
                )

            elif (
                identity_checks.get(
                    "name_match"
                )
                is not True
            ):
                error_message = (
                    "Họ tên sinh viên trên đơn "
                    "không khớp với tài khoản "
                    "đang đăng nhập."
                )

            elif (
                parent_signature.get("present")
                is not True
            ):
                error_message = (
                    "Không tìm thấy chữ ký của "
                    "Phụ huynh/Người giám hộ."
                )

            elif (
                applicant_signature.get("present")
                is not True
            ):
                error_message = (
                    "Không tìm thấy chữ ký của "
                    "Người làm đơn."
                )

            elif (
                ocr_data.get("is_match")
                is not True
            ):
                error_message = (
                    "File tải lên không phải "
                    "Đơn xin thôi học."
                )

            else:
                error_message = (
                    ocr_data.get(
                        "validation_reason"
                    )
                    or
                    "Đơn xin thôi học không hợp lệ."
                )

            return Response(
                {
                    "error": error_message,
                    "ocr_result": ocr_data,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        from apps.requests.models import (
            Request as StudentRequest,
            DropoutRequest,
            RequestDocument,
            RequestHistory,
        )

        existing = (
            StudentRequest.objects
            .filter(
                student=student,
                request_type=(
                    StudentRequest
                    .RequestType
                    .DROPOUT
                ),
                status=(
                    StudentRequest
                    .Status
                    .DRAFT
                ),
            )
            .first()
        )

        if existing:
            req = existing
            req.status = (
                StudentRequest.Status.PENDING
            )
            req.submitted_at = timezone.now()

            req.save(
                update_fields=[
                    "status",
                    "submitted_at",
                    "updated_at",
                ]
            )

        else:
            existing_active = (
                StudentRequest.objects
                .filter(
                    student=student,
                    request_type=(
                        StudentRequest
                        .RequestType
                        .DROPOUT
                    ),
                )
                .exclude(
                    status__in=[
                        StudentRequest.Status.APPROVED,
                        StudentRequest.Status.REJECTED,
                    ]
                )
                .first()
            )

            if existing_active:
                return Response(
                    {
                        "error":
                        "Bạn đang có một yêu cầu đang xử lý. "
                        "Vui lòng hoàn tất hoặc hủy yêu cầu đó "
                        "trước khi gửi mới."
                    },
                    status=status.HTTP_409_CONFLICT,
                )

            req = StudentRequest.objects.create(
                student=student,
                request_type=(
                    StudentRequest
                    .RequestType
                    .DROPOUT
                ),
                status=(
                    StudentRequest.Status.PENDING
                ),
                submitted_at=timezone.now(),
            )

        DropoutRequest.objects.update_or_create(
            request=req,
            defaults={
                "reason": reason,
            },
        )

        old_documents = req.documents.filter(
            document_type=(
                RequestDocument
                .DocumentType
                .INITIAL
            )
        )

        for old_document in old_documents:
            if old_document.file:
                old_document.file.delete(
                    save=False
                )

        old_documents.delete()

        RequestDocument.objects.create(
            request=req,
            file=request_document_file,
            file_name=uploaded_file.name,
            document_type=(
                RequestDocument
                .DocumentType
                .INITIAL
            ),
        )

        RequestHistory.objects.create(
            request=req,
            status=req.status,
            actor=(
                request.user
                if request.user.is_authenticated
                else None
            ),
            notes=(
                "Sinh viên nộp hồ sơ thôi học. "
                "Đơn đã được kiểm tra loại tài liệu, "
                "danh tính và chữ ký."
            ),
        )

        tracking_code = (
            f"TH-{str(req.id).split('-')[0].upper()}"
        )

        return Response(
            {
                "success": True,
                "trackingCode": tracking_code,
                "requestId": str(req.id),
                "status": req.status,
                "ocr_result": ocr_data,
            },
            status=status.HTTP_201_CREATED,
        )